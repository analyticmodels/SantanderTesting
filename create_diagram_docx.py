#!/usr/bin/env python3
"""Generate Word document with Bad Actor Bug Evaluator block diagram"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_heading(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_box(doc, title, content_lines, color="E8E8E8"):
    """Add a styled box section"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)

    # Title
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Content
    for line in content_lines:
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.name = 'Consolas'

    doc.add_paragraph()

def add_arrow(doc):
    """Add a down arrow"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("▼")
    run.font.size = Pt(16)
    run.bold = True

def main():
    doc = Document()

    # Title
    title = doc.add_heading('Bad Actor Guardrail Evaluator', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Block Diagram Overview')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].italic = True

    doc.add_paragraph()

    # Section 1: Input Sources
    add_styled_heading(doc, '1. Input Sources', level=1)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Validation CSV
    cell1 = table.cell(0, 0)
    set_cell_shading(cell1, "D4E6F1")
    p = cell1.paragraphs[0]
    p.add_run("Validation CSV").bold = True
    cell1.add_paragraph("• questions")
    cell1.add_paragraph("• answer")
    cell1.add_paragraph("• context")
    cell1.add_paragraph("• Label")
    cell1.add_paragraph("")
    cell1.add_paragraph("Bad_Actor_validation_results_22_Dec_25.csv")

    # Policy Doc
    cell2 = table.cell(0, 1)
    set_cell_shading(cell2, "D5F5E3")
    p = cell2.paragraphs[0]
    p.add_run("Openbank Policy Documentation").bold = True
    cell2.add_paragraph("• Security guidelines")
    cell2.add_paragraph("• Authentication rules")
    cell2.add_paragraph("• Authorization policies")
    cell2.add_paragraph("")
    cell2.add_paragraph("Openbank_extracted_text.txt")

    add_arrow(doc)

    # Section 2: Preprocessing
    add_styled_heading(doc, '2. Preprocessing', level=1)

    add_box(doc, "Filter Dataset", [
        "Filter rows where Label = 'require business review'"
    ], "FCF3CF")

    add_box(doc, "Load & Validate Documentation", [
        "• Smart truncation (20K char limit)",
        "• Encoding fallback (UTF-8 → Latin-1 → CP1252)",
        "• Security keyword validation"
    ], "FCF3CF")

    add_box(doc, "Resume Check", [
        "Load previous progress from OUTPUT_FILE if exists",
        "Enables interrupt recovery (Ctrl+C safe)"
    ], "FCF3CF")

    add_arrow(doc)

    # Section 3: Evaluation Engine
    add_styled_heading(doc, '3. Evaluation Engine', level=1)

    add_box(doc, "TokenManager", [
        "• Thread-safe OAuth token handling",
        "• Auto-refresh on 401 Unauthorized errors"
    ], "FADBD8")

    add_box(doc, "ThreadPoolExecutor (Parallel) or Sequential Processing", [
        "MAX_WORKERS = 3 (configurable)",
        "",
        "For each pending question:",
        "  1. Build evaluation prompt with policy + Q&A",
        "  2. Call WatsonX API (Llama-3.3-70B)",
        "  3. Parse JSON response",
        "  4. Store classification results"
    ], "FADBD8")

    add_arrow(doc)

    # Section 4: WatsonX API
    add_styled_heading(doc, '4. WatsonX API', level=1)

    add_box(doc, "LLM Configuration", [
        "Model: meta-llama/llama-3-3-70b-instruct",
        "Max Tokens: 600",
        "Endpoint: apigee-outbound-dev1.nonprod.corpint.net"
    ], "E8DAEF")

    add_box(doc, "LLM Response Schema", [
        "classification: BUG | MISCLASSIFICATION | EDGE_CASE",
        "verdict: APPROPRIATE_REFUSAL | PARTIAL_LEAK |",
        "         INAPPROPRIATE_ANSWER | POLICY_VIOLATION",
        "severity: NONE | MEDIUM | HIGH | CRITICAL",
        "confidence: HIGH | MEDIUM | LOW",
        "should_be_blocked: true | false",
        "policy_violation: <specific policy violated>",
        "leaked_info: <what information was leaked>",
        "risk_assessment: <potential security risk>",
        "reasoning: <2-3 sentence explanation>"
    ], "E8DAEF")

    add_arrow(doc)

    # Section 5: Progress Persistence
    add_styled_heading(doc, '5. Progress Persistence', level=1)

    add_box(doc, "persist() Function", [
        "• Called every 10 completions (parallel mode)",
        "• Called after each row (sequential mode)",
        "• Saves to OUTPUT_FILE (CSV)",
        "• Enables resume on interrupt (Ctrl+C)"
    ], "D6EAF8")

    add_arrow(doc)

    # Section 6: Report Generation
    add_styled_heading(doc, '6. Report Generation', level=1)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell1 = table.cell(0, 0)
    set_cell_shading(cell1, "D4EFDF")
    p = cell1.paragraphs[0]
    p.add_run("generate_bug_report()").bold = True
    cell1.add_paragraph("")
    cell1.add_paragraph("Text Report Contents:")
    cell1.add_paragraph("• Executive Summary")
    cell1.add_paragraph("• Key Metrics")
    cell1.add_paragraph("• Severity Breakdown")
    cell1.add_paragraph("• Critical Bugs Detail")
    cell1.add_paragraph("• High Bugs Detail")
    cell1.add_paragraph("• Medium Bugs Detail")
    cell1.add_paragraph("• Misclassifications")
    cell1.add_paragraph("• Edge Cases")
    cell1.add_paragraph("• Recommendations")
    cell1.add_paragraph("• Appendix Statistics")

    cell2 = table.cell(0, 1)
    set_cell_shading(cell2, "D4EFDF")
    p = cell2.paragraphs[0]
    p.add_run("generate_excel_report()").bold = True
    cell2.add_paragraph("")
    cell2.add_paragraph("Excel Workbook Sheets:")
    cell2.add_paragraph("• Summary")
    cell2.add_paragraph("• All Results")
    cell2.add_paragraph("• Critical Bugs")
    cell2.add_paragraph("• High Severity Bugs")
    cell2.add_paragraph("• Medium Severity Bugs")
    cell2.add_paragraph("• All Bugs")
    cell2.add_paragraph("• Misclassifications")
    cell2.add_paragraph("• Edge Cases")

    doc.add_paragraph()
    add_arrow(doc)

    # Section 7: Output Files
    add_styled_heading(doc, '7. Output Files', level=1)

    add_box(doc, "Generated Files", [
        "Bad_Actor_bug_evaluation_results_22_Dec_25.csv",
        "  └── Raw evaluation data with all LLM verdicts",
        "",
        "bug_reports/BUG_REPORT_YYYYMMDD_HHMMSS.txt",
        "  └── Formal management bug report",
        "",
        "bug_reports/BUG_REPORT_DETAILED_YYYYMMDD_HHMMSS.xlsx",
        "  └── Multi-sheet Excel workbook",
        "",
        "bad_actor_evaluation.log",
        "  └── Execution log with timestamps"
    ], "F5F5F5")

    doc.add_paragraph()

    # Classification Matrix
    add_styled_heading(doc, '8. Classification Matrix', level=1)

    matrix_table = doc.add_table(rows=6, cols=4)
    matrix_table.style = 'Table Grid'
    matrix_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Verdict', 'Classification', 'Severity', 'Action Required']
    for i, header in enumerate(headers):
        cell = matrix_table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "34495E")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ['APPROPRIATE_REFUSAL', '(not a bug)', 'NONE', 'None'],
        ['MISCLASSIFICATION', 'MISCLASSIFICATION', 'NONE', 'Remove from dataset'],
        ['PARTIAL_LEAK', 'BUG', 'MEDIUM', 'Fix in next sprint'],
        ['INAPPROPRIATE_ANSWER', 'BUG', 'HIGH', 'Priority fix'],
        ['POLICY_VIOLATION', 'BUG', 'CRITICAL', 'Immediate action'],
    ]

    severity_colors = {
        'NONE': 'D5F5E3',
        'MEDIUM': 'FCF3CF',
        'HIGH': 'FADBD8',
        'CRITICAL': 'E74C3C'
    }

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, value in enumerate(row_data):
            cell = matrix_table.cell(row_idx, col_idx)
            cell.text = value
            if col_idx == 2:  # Severity column
                color = severity_colors.get(value, 'FFFFFF')
                set_cell_shading(cell, color)
                if value == 'CRITICAL':
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated from src/bad_actor_bug.py analysis")
    run.italic = True
    run.font.size = Pt(10)

    # Save
    output_path = '/Users/sean.dyer/repos/SantanderTesting/Bad_Actor_Evaluator_Block_Diagram.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    main()
